import re
import markdown
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

def add_custom_style(doc):
    """添加自定义样式"""
    # 标题样式
    styles = doc.styles
    
    # 一级标题样式
    if 'Heading 1' in styles:
        h1_style = styles['Heading 1']
        h1_style.font.name = '宋体'
        h1_style.font.size = Pt(16)
        h1_style.font.bold = True
    
    # 二级标题样式
    if 'Heading 2' in styles:
        h2_style = styles['Heading 2']
        h2_style.font.name = '宋体'
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True
    
    # 三级标题样式
    if 'Heading 3' in styles:
        h3_style = styles['Heading 3']
        h3_style.font.name = '宋体'
        h3_style.font.size = Pt(12)
        h3_style.font.bold = True
    
    # 正文样式
    if 'Normal' in styles:
        normal_style = styles['Normal']
        normal_style.font.name = '宋体'
        normal_style.font.size = Pt(10.5)

def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.15):
    """设置段落间距"""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(before)
    paragraph_format.space_after = Pt(after)
    paragraph_format.line_spacing = line_spacing

def markdown_to_word(md_file_path, word_file_path):
    """将Markdown文件转换为Word文档"""
    # 读取Markdown文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 创建Word文档
    doc = Document()
    add_custom_style(doc)
    
    # 分割Markdown内容为行
    lines = md_content.split('\n')
    
    current_paragraph = None
    in_table = False
    table_data = []
    
    for line in lines:
        # 处理标题
        if line.startswith('# '):
            if current_paragraph:
                set_paragraph_spacing(current_paragraph, after=6)
            heading = doc.add_heading(line[2:], level=1)
            set_paragraph_spacing(heading, before=12, after=6)
            current_paragraph = None
        elif line.startswith('## '):
            if current_paragraph:
                set_paragraph_spacing(current_paragraph, after=6)
            heading = doc.add_heading(line[3:], level=2)
            set_paragraph_spacing(heading, before=12, after=6)
            current_paragraph = None
        elif line.startswith('### '):
            if current_paragraph:
                set_paragraph_spacing(current_paragraph, after=6)
            heading = doc.add_heading(line[4:], level=3)
            set_paragraph_spacing(heading, before=12, after=6)
            current_paragraph = None
        # 处理表格
        elif line.startswith('|') and line.endswith('|'):
            # 表格行
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(cells)
            in_table = True
        elif in_table and line.startswith('|'):
            # 表格分隔符，跳过
            continue
        elif in_table and not line.strip():
            # 表格结束，创建表格
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row_data in enumerate(table_data):
                    for j, cell_data in enumerate(row_data):
                        table.cell(i, j).text = cell_data
                        table.cell(i, j).paragraphs[0].runs[0].font.name = '宋体'
                        table.cell(i, j).paragraphs[0].runs[0].font.size = Pt(10.5)
                set_paragraph_spacing(doc.paragraphs[-1], after=12)
                table_data = []
                in_table = False
                current_paragraph = None
        # 处理空行
        elif not line.strip():
            if current_paragraph:
                set_paragraph_spacing(current_paragraph, after=6)
                current_paragraph = None
            doc.add_paragraph()
        # 处理普通段落
        else:
            # 如果是列表项
            if line.strip().startswith('- ') or line.strip().startswith('* '):
                if current_paragraph:
                    set_paragraph_spacing(current_paragraph, after=6)
                p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
                set_paragraph_spacing(p, after=3)
                current_paragraph = None
            # 如果是数字列表
            elif re.match(r'^\d+\.', line.strip()):
                if current_paragraph:
                    set_paragraph_spacing(current_paragraph, after=6)
                p = doc.add_paragraph(line.strip(), style='List Number')
                set_paragraph_spacing(p, after=3)
                current_paragraph = None
            # 普通段落
            else:
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph()
                    current_paragraph.add_run(line)
                    current_paragraph.runs[0].font.name = '宋体'
                    current_paragraph.runs[0].font.size = Pt(10.5)
                else:
                    current_paragraph.add_run('\n' + line)
    
    # 处理最后一个表格（如果文档以表格结束）
    if in_table and table_data:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Table Grid'
        for i, row_data in enumerate(table_data):
            for j, cell_data in enumerate(row_data):
                table.cell(i, j).text = cell_data
                table.cell(i, j).paragraphs[0].runs[0].font.name = '宋体'
                table.cell(i, j).paragraphs[0].runs[0].font.size = Pt(10.5)
    
    # 保存Word文档
    doc.save(word_file_path)
    print(f"Word文档已保存至: {word_file_path}")

if __name__ == "__main__":
    md_file = "中期报告.md"
    word_file = "中期报告.docx"
    markdown_to_word(md_file, word_file)